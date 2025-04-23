from flask import Flask, render_template, request, jsonify
from docx import Document
import re
import os
from werkzeug.utils import secure_filename

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['ALLOWED_EXTENSIONS'] = {'docx'}

# Asegurarnos de que la carpeta de uploads existe
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def extraer_marcadores(doc_path):
    doc = Document(doc_path)
    marcadores = []
    
    for paragraph in doc.paragraphs:
        text = paragraph.text
        matches = re.findall(r'\+(.*?)\+', text)
        for match in matches:
            marcador = match.strip()
            if marcador not in marcadores:
                marcadores.append(marcador)
    
    return marcadores

def generar_documento(template_path, reemplazos, output_filename):
    doc = Document(template_path)
    
    for paragraph in doc.paragraphs:
        runs = paragraph.runs
        full_text = ''.join([run.text for run in runs])
        
        if '+' in full_text:
            new_runs = []
            for run in runs:
                text = run.text
                if '+' in text:
                    for marcador, reemplazo in reemplazos.items():
                        text = text.replace(f'+{marcador}+', reemplazo)
                new_runs.append((text, run))
            
            paragraph.clear()
            for text, original_run in new_runs:
                new_run = paragraph.add_run(text)
                new_run.bold = original_run.bold
                new_run.italic = original_run.italic
                new_run.underline = original_run.underline
                new_run.font.name = original_run.font.name
                new_run.font.size = original_run.font.size
                if original_run.font.color.rgb:
                    new_run.font.color.rgb = original_run.font.color.rgb
    
    output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
    doc.save(output_path)
    return output_path

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        # Manejar la subida del archivo
        if 'file' not in request.files:
            return jsonify({'error': 'No file part'})
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No selected file'})
        
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            template_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(template_path)
            
            marcadores = extraer_marcadores(template_path)
            return jsonify({'marcadores': marcadores})
        
        return jsonify({'error': 'Invalid file type'})
    
    return render_template('index.html')

@app.route('/preview', methods=['POST'])
def preview():
    data = request.json
    template_path = os.path.join(app.config['UPLOAD_FOLDER'], data['filename'])
    reemplazos = data['reemplazos']
    
    # Crear una versión preview sin guardar el documento
    doc = Document(template_path)
    preview_text = []
    
    for paragraph in doc.paragraphs:
        text = paragraph.text
        for marcador, reemplazo in reemplazos.items():
            text = text.replace(f'+{marcador}+', reemplazo)
        preview_text.append(text)
    
    return jsonify({'preview': '\n'.join(preview_text)})

@app.route('/generate', methods=['POST'])
def generate():
    data = request.json
    template_path = os.path.join(app.config['UPLOAD_FOLDER'], data['filename'])
    reemplazos = data['reemplazos']
    output_filename = secure_filename(data['output_filename'])
    
    if not output_filename.endswith('.docx'):
        output_filename += '.docx'
    
    try:
        output_path = generar_documento(template_path, reemplazos, output_filename)
        return jsonify({
            'success': True,
            'message': f'Documento generado exitosamente: {output_filename}',
            'filename': output_filename
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Error al generar el documento: {str(e)}'
        })

if __name__ == '__main__':
    app.run(debug=True)