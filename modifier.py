from docx import Document
import re

def extraer_marcadores(doc_path):
    doc = Document(doc_path)
    marcadores = set()
    
    for paragraph in doc.paragraphs:
        text = paragraph.text
        # Buscamos todos los patrones entre ++ sin importar formato
        matches = re.findall(r'\+(.*?)\+', text)
        for match in matches:
            marcadores.add(match.strip())
    
    return sorted(marcadores)

def reemplazar_marcadores(doc_path, reemplazos, nombre_salida):
    doc = Document(doc_path)
    
    for paragraph in doc.paragraphs:
        # Necesitamos procesar los runs para mantener el formato
        runs = paragraph.runs
        full_text = ''.join([run.text for run in runs])
        
        if '+' in full_text:  # Solo procesar párrafos con marcadores
            # Reconstruimos el párrafo con los reemplazos
            new_runs = []
            for run in runs:
                text = run.text
                if '+' in text:
                    # Reemplazamos cada marcador manteniendo el formato original
                    for marcador, reemplazo in reemplazos.items():
                        text = text.replace(f'+{marcador}+', reemplazo)
                new_runs.append((text, run))
            
            # Limpiamos el párrafo y agregamos los runs con formato
            paragraph.clear()
            for text, original_run in new_runs:
                new_run = paragraph.add_run(text)
                # Copiamos todo el formato del run original
                new_run.bold = original_run.bold
                new_run.italic = original_run.italic
                new_run.underline = original_run.underline
                new_run.font.name = original_run.font.name
                new_run.font.size = original_run.font.size
                new_run.font.color.rgb = original_run.font.color.rgb
    
    doc.save(nombre_salida)
    print(f"\nDocumento generado exitosamente: {nombre_salida}")

def main():
    template_path = 'template.docx'
    
    print("Procesando el documento template...")
    marcadores = extraer_marcadores(template_path)
    
    if not marcadores:
        print("No se encontraron marcadores (texto entre +) en el documento.")
        return
    
    print("\nSe encontraron los siguientes marcadores para reemplazar:")
    reemplazos = {}
    
    for marcador in marcadores:
        reemplazo = input(f"Reemplazar {marcador} por: ")
        reemplazos[marcador] = reemplazo
    
    nombre_archivo = input("\nIngrese el nombre del archivo de salida (sin extensión .docx): ")
    nombre_salida = f"{nombre_archivo.strip()}.docx"
    
    reemplazar_marcadores(template_path, reemplazos, nombre_salida)

if __name__ == "__main__":
    print("=== Generador de documentos desde template ===")
    print("Asegúrese de tener un archivo 'template.docx' en la misma carpeta")
    main()